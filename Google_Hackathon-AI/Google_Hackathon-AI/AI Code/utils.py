import logging
import fitz  # PyMuPDF
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
import os

# Try to import optional dependencies
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    HuggingFaceEmbeddings = None
    logging.warning("HuggingFace embeddings not available")

try:
    from langchain_core.documents import Document
except ImportError:
    # Fallback Document class if langchain_core not available
    class Document:
        def __init__(self, page_content: str, metadata: Dict = None):
            self.page_content = page_content
            self.metadata = metadata or {}

logger = logging.getLogger(__name__)


def validate_file_type(file_path: str) -> bool:
    """Validate if the file type is supported."""
    supported_extensions = {'.pdf', '.doc', '.docx', '.txt'}
    file_extension = Path(file_path).suffix.lower()
    return file_extension in supported_extensions


def create_metadata(chunk, file_path: str) -> Dict[str, Any]:
    """Create metadata for document chunks."""
    return {
        'source': file_path,
        'file_name': Path(file_path).name,
        'file_type': Path(file_path).suffix.lower(),
        'processed_at': datetime.now().isoformat(),
        'chunk_index': getattr(chunk, 'chunk_index', 0)
    }


class DocumentProcessor:
    """Handles document processing for multiple file types with improved error handling."""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentProcessor")

    def load_document(self, file_path: str) -> List[Document]:
        """Load document and return as LangChain Document objects."""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
                
            file_extension = Path(file_path).suffix.lower()
            self.logger.info(f"Loading {file_extension} file: {file_path}")
            
            if file_extension == '.pdf':
                return self._load_pdf(file_path)
            elif file_extension == '.txt':
                return self._load_text(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._load_word(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error loading document {file_path}: {e}")
            raise
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Load PDF document with improved text extraction."""
        documents = []
        try:
            with fitz.open(file_path) as doc:
                total_pages = doc.page_count
                self.logger.info(f"Processing PDF with {total_pages} pages")
                
                for page_num in range(total_pages):
                    try:
                        page = doc.load_page(page_num)
                        text = page.get_text()
                        
                        # If no text extracted, try alternative methods
                        if not text.strip():
                            # Try to extract from images or other elements
                            text = page.get_text("text")  # alternative extraction
                        
                        if text.strip():  # Only add non-empty pages
                            documents.append(Document(
                                page_content=text,
                                metadata={
                                    'source': file_path,
                                    'page': page_num + 1,  # 1-indexed page numbers
                                    'total_pages': total_pages,
                                    'file_type': 'pdf',
                                    'file_size': os.path.getsize(file_path)
                                }
                            ))
                        else:
                            self.logger.debug(f"Page {page_num + 1} contains no extractable text")
                            
                    except Exception as e:
                        self.logger.warning(f"Error processing page {page_num + 1}: {e}")
                        continue
                        
                if not documents:
                    self.logger.warning(f"No text content extracted from PDF: {file_path}")
                    # Return a document indicating the PDF was processed but contained no text
                    documents.append(Document(
                        page_content="No extractable text found in this PDF document.",
                        metadata={
                            'source': file_path,
                            'file_type': 'pdf',
                            'total_pages': total_pages,
                            'error': 'No extractable text'
                        }
                    ))
                    
        except Exception as e:
            self.logger.error(f"Error loading PDF {file_path}: {e}")
            raise
        return documents
    
    def _load_text(self, file_path: str) -> List[Document]:
        """Load text document with improved encoding handling."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        self.logger.info(f"Successfully read text file with {encoding} encoding")
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode text file with any supported encoding")
                
            return [Document(
                page_content=content,
                metadata={
                    'source': file_path,
                    'file_type': 'txt',
                    'file_size': os.path.getsize(file_path)
                }
            )]
        except Exception as e:
            self.logger.error(f"Error loading text file {file_path}: {e}")
            raise
    
    def _load_word(self, file_path: str) -> List[Document]:
        """Load Word document with improved error handling."""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            # Try to use python-docx for .docx files
            if file_extension == '.docx':
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    
                    if not text.strip():
                        # Try to extract from tables if no paragraph text
                        table_text = []
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        table_text.append(cell.text.strip())
                        text = '\n'.join(table_text)
                    
                    return [Document(
                        page_content=text,
                        metadata={
                            'source': file_path,
                            'file_type': 'docx',
                            'file_size': os.path.getsize(file_path)
                        }
                    )]
                except ImportError:
                    self.logger.warning("python-docx not available for .docx files")
                    return [Document(
                        page_content="Error: Cannot process .docx file. python-docx not installed.",
                        metadata={
                            'source': file_path,
                            'file_type': 'docx',
                            'error': 'python-docx not available'
                        }
                    )]
            
            # For .doc files or fallback
            elif file_extension == '.doc':
                # .doc files require more complex handling
                self.logger.warning("Legacy .doc files not fully supported. Consider converting to .docx")
                return [Document(
                    page_content="Error: Legacy .doc files are not supported. Please convert to .docx format.",
                    metadata={
                        'source': file_path,
                        'file_type': 'doc',
                        'error': 'Legacy format not supported'
                    }
                )]
            
            else:
                raise ValueError(f"Unsupported Word format: {file_extension}")
                
        except Exception as e:
            self.logger.error(f"Error loading Word document {file_path}: {e}")
            raise

    def get_page_count(self, pdf_path: str) -> int:
        """Gets the total number of pages in a PDF document."""
        try:
            with fitz.open(pdf_path) as doc:
                return doc.page_count
        except Exception as e:
            logger.error(f"Error getting page count for {pdf_path}: {e}")
            raise

    def get_full_text(self, pdf_path: str) -> str:
        """Extracts and concatenates text from all pages of a PDF."""
        try:
            with fitz.open(pdf_path) as doc:
                # Using a generator expression for memory efficiency
                return "\n\n".join(page.get_text() for page in doc)
        except Exception as e:
            logger.error(f"Error extracting full text from {pdf_path}: {e}")
            raise

    def get_text_from_page(self, pdf_path: str, page_number: int) -> str:
        """Extracts text from a specific page of a PDF."""
        try:
            with fitz.open(pdf_path) as doc:
                if not 0 <= page_number < doc.page_count:
                    raise ValueError("Page number out of bounds.")
                page = doc.load_page(page_number)
                return page.get_text()
        except Exception as e:
            logger.error(f"Error extracting text from page {page_number} of {pdf_path}: {e}")
            raise


class EmbeddingManager:
    """Manages embedding models with better error handling."""

    def __init__(self, model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        """
        Initializes the EmbeddingManager.
        Args:
            model_name: The name of the sentence-transformer model to use.
        """
        self.model_name = model_name
        self._embeddings = None
        self.logger = logging.getLogger(f"{__name__}.EmbeddingManager")

    def get_embeddings(self):
        """Lazily loads and returns the embedding model instance."""
        if self._embeddings is None:
            if HuggingFaceEmbeddings is None:
                raise ImportError("HuggingFace embeddings not available. Install with: pip install langchain-huggingface")
            
            try:
                self._embeddings = HuggingFaceEmbeddings(
                    model_name=self.model_name,
                    model_kwargs={"device": "cpu"},
                    encode_kwargs={"normalize_embeddings": True},
                )
                self.logger.info(f"Loaded embedding model: {self.model_name}")
            except Exception as e:
                self.logger.error(f"Failed to load embedding model {self.model_name}: {e}")
                raise
                
        return self._embeddings

    def is_available(self) -> bool:
        """Check if embeddings are available."""
        return HuggingFaceEmbeddings is not None
